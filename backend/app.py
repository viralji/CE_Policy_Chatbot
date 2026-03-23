import os
import re
import time
import traceback
from collections import OrderedDict
from flask import Flask, request, jsonify, send_from_directory, g
from flask_cors import CORS
from dotenv import load_dotenv
from PyPDF2 import PdfReader
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX files will be skipped. Run: pip install python-docx", flush=True)

# ---------------------- Load API key and config from .env ----------------------
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
BACKEND_PORT = int(os.getenv("PORT", "4001"))
BASE_URL = os.getenv("BASE_URL", f"http://localhost:{BACKEND_PORT}")
MAX_QUESTION_LENGTH = 2000

if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY or GEMINI_API_KEY not found in .env file")

os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY

# ---------------------- Flask App Setup ----------------------
app = Flask(__name__)
CORS(app)

# Resolve data folder relative to this file so the app works from any working directory
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, 'data')
FAISS_INDEX_DIR = os.path.join(BASE_DIR, 'faiss_index')
FAISS_INDEX_FILE = os.path.join(FAISS_INDEX_DIR, 'index.faiss')

# ---------------------- Load PDFs and DOCX with metadata ----------------------
def load_all_docs_with_metadata(folder=DATA_DIR):
    docs = []
    if not os.path.isdir(folder):
        print(f"Warning: data folder not found at {folder}", flush=True)
        return docs

    for file in sorted(os.listdir(folder)):
        filepath = os.path.join(folder, file)
        if file.endswith('.pdf'):
            try:
                reader = PdfReader(filepath)
                for i, page in enumerate(reader.pages):
                    content = page.extract_text()
                    if content and content.strip():
                        docs.append({
                            'text': content,
                            'metadata': {'source': file, 'page': i + 1}
                        })
            except Exception as e:
                print(f"Warning: failed to load PDF '{file}': {e}", flush=True)
        elif file.endswith('.docx'):
            if not DOCX_AVAILABLE:
                print(f"Warning: skipping DOCX '{file}' (python-docx not installed)", flush=True)
                continue
            try:
                doc = DocxDocument(filepath)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                # Also extract table text
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                paragraphs.append(cell_text)
                text = '\n'.join(paragraphs)
                if text:
                    docs.append({
                        'text': text,
                        'metadata': {'source': file, 'page': None}
                    })
            except Exception as e:
                print(f"Warning: failed to load DOCX '{file}': {e}", flush=True)
    return docs


def index_is_stale(folder=DATA_DIR):
    """Return True if any source file is newer than the saved FAISS index, or if embedding model changed."""
    model_file = os.path.join(FAISS_INDEX_DIR, ".embedding_model")
    if not os.path.exists(FAISS_INDEX_FILE):
        return True
    # Rebuild if embedding model changed (or unknown model from before we tracked it)
    if not os.path.exists(model_file):
        print("No .embedding_model file, rebuilding index...", flush=True)
        return True
    with open(model_file) as f:
        if f.read().strip() != EMBEDDING_MODEL:
            print(f"Embedding model changed, rebuilding index...", flush=True)
            return True
    index_mtime = os.path.getmtime(FAISS_INDEX_FILE)
    for file in os.listdir(folder):
        if file.endswith(('.pdf', '.docx')):
            if os.path.getmtime(os.path.join(folder, file)) > index_mtime:
                return True
    return False


print("Loading documents...", flush=True)
docs = load_all_docs_with_metadata()
total_text_length = sum(len(doc['text']) for doc in docs)
print(f"Loaded {len(docs)} pages/documents, {total_text_length} characters total", flush=True)

text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
chunks = []
for doc in docs:
    splits = text_splitter.split_text(doc['text'])
    for split in splits:
        chunks.append({'text': split, 'metadata': doc['metadata']})
print(f"Total chunks: {len(chunks)}", flush=True)

# ---------------------- Create Embeddings and Vectorstore ----------------------
# models/gemini-embedding-001 is the supported model in v1beta API used by langchain-google-genai
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "models/gemini-embedding-001")
LLM_MODEL = os.getenv("LLM_MODEL", "models/gemini-2.5-pro")
print(f"Creating/loading FAISS index (embedding: {EMBEDDING_MODEL})...", flush=True)
embeddings = GoogleGenerativeAIEmbeddings(model=EMBEDDING_MODEL)

if index_is_stale():
    print("Building FAISS index (new/changed documents detected)...", flush=True)
    texts = [chunk['text'] for chunk in chunks]
    metadatas = [chunk['metadata'] for chunk in chunks]
    # Retry the embedding API call at startup — a transient 500 must not crash the server.
    _build_retries = 5
    for _attempt in range(_build_retries):
        try:
            vectorstore = FAISS.from_texts(texts, embedding=embeddings, metadatas=metadatas)
            break
        except Exception as _e:
            _err = str(_e).lower()
            if any(kw in _err for kw in ("500", "internal", "embedding", "429", "quota", "rate", "overloaded")) \
                    and _attempt < _build_retries - 1:
                _wait = 2 * (2 ** _attempt)
                print(f"Embedding API error during index build (attempt {_attempt + 1}/{_build_retries}), "
                      f"retrying in {_wait}s: {_e}", flush=True)
                time.sleep(_wait)
            else:
                raise
    vectorstore.save_local(FAISS_INDEX_DIR)
    with open(os.path.join(FAISS_INDEX_DIR, ".embedding_model"), "w") as f:
        f.write(EMBEDDING_MODEL)
    print("FAISS index built and saved.", flush=True)
else:
    print("Loading FAISS index from disk...", flush=True)
    vectorstore = FAISS.load_local(FAISS_INDEX_DIR, embeddings, allow_dangerous_deserialization=True)

print("Vectorstore ready!", flush=True)

retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
llm = ChatGoogleGenerativeAI(model=LLM_MODEL, temperature=0)
print(f"LLM ready ({LLM_MODEL})!", flush=True)

# ---------------------- System prompt for RAG ----------------------
_SYSTEM_PROMPT = (
    "You are CloudExtel's HR policy assistant. Answer questions using only the "
    "company policy information provided below. Be specific and cite relevant details.\n\n"
    "If the answer is not found in the provided context, say: "
    "\"I don't have information about that in the company policies. Please contact HR directly.\"\n\n"
    "Company Policy Context:\n{context}"
)

# ---------------------- Per-user conversation history ----------------------
# Each user gets a list of {human, ai} dicts. LRU eviction keeps memory bounded.
_user_histories: OrderedDict = OrderedDict()
MAX_USER_HISTORIES = 200
MAX_HISTORY_TURNS = 10   # Keep last 10 Q&A pairs per user


def get_user_history(user_email: str) -> list:
    """Return (and lazily create) per-user conversation history with LRU eviction."""
    if user_email in _user_histories:
        _user_histories.move_to_end(user_email)
        return _user_histories[user_email]
    if len(_user_histories) >= MAX_USER_HISTORIES:
        evicted = next(iter(_user_histories))
        del _user_histories[evicted]
        print(f"Evicted oldest user history: {evicted}", flush=True)
    _user_histories[user_email] = []
    return _user_histories[user_email]


def _retry(fn, *args, label="API call"):
    """Call fn(*args) with up to 5 retries on transient Google errors."""
    _RETRYABLE = ("500", "internal", "embedding", "429", "quota", "rate", "overloaded")
    max_retries = 5
    for attempt in range(max_retries):
        try:
            return fn(*args)
        except Exception as e:
            err_str = str(e).lower()
            if any(kw in err_str for kw in _RETRYABLE) and attempt < max_retries - 1:
                wait = 2 * (2 ** attempt)   # 2, 4, 8, 16 s
                print(f"{label} error (attempt {attempt + 1}/{max_retries}), retrying in {wait}s: {e}", flush=True)
                time.sleep(wait)
            else:
                raise


# ---------------------- Auth (Microsoft Azure AD) ----------------------
from auth_middleware import require_auth

# ---------------------- API Endpoints ----------------------
@app.route('/files/<path:filename>')
def serve_file(filename):
    return send_from_directory(DATA_DIR, filename)

@app.route('/api/auth/me', methods=['GET'])
@require_auth
def auth_me():
    return jsonify(g.user)

@app.route('/api/chat', methods=['POST'])
@require_auth
def chat():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Request body must be JSON'}), 400
        question = data.get('question', '').strip()
        if not question:
            return jsonify({'error': 'Question is required'}), 400
        if len(question) > MAX_QUESTION_LENGTH:
            return jsonify({'error': f'Question too long (max {MAX_QUESTION_LENGTH} characters)'}), 400

        user_email = g.user.get('email', 'anonymous')
        history = get_user_history(user_email)

        # Step 1: Retrieve relevant docs (embedding call — isolated for independent retry)
        source_docs = _retry(retriever.invoke, question, label="Embedding/retrieval")

        # Step 2: Build prompt — history + retrieved context + question
        context = "\n\n".join(doc.page_content for doc in source_docs)
        messages = [SystemMessage(content=_SYSTEM_PROMPT.format(context=context))]
        for turn in history[-MAX_HISTORY_TURNS:]:
            messages.append(HumanMessage(content=turn['human']))
            messages.append(AIMessage(content=turn['ai']))
        messages.append(HumanMessage(content=question))

        # Step 3: LLM call (independent retry from embedding)
        response = _retry(llm.invoke, messages, label="LLM")
        answer = response.content

        # Persist this exchange in history
        history.append({'human': question, 'ai': answer})

        # Format answer: strip leading markdown markers, keep non-empty lines
        answer_lines = []
        for line in answer.split('\n'):
            stripped = re.sub(r'^[-*•#]+\s*', '', line).rstrip()
            if stripped:
                answer_lines.append(stripped)

        # Deduplicate sources: unique (file, page) pairs, preserving order
        seen = set()
        sources = []
        for doc in source_docs:
            meta = doc.metadata if hasattr(doc, 'metadata') else doc.get('metadata', {})
            file = meta.get('source')
            page = meta.get('page')
            key = (file, page)
            if file and key not in seen:
                seen.add(key)
                link = f"{BASE_URL}/files/{file}#page={page}" if page else f"{BASE_URL}/files/{file}"
                sources.append({'file': file, 'page': page, 'link': link})

        return jsonify({
            'response': answer_lines,
            'sources': sources
        })
    except Exception as e:
        raw = str(e) or f"{type(e).__name__}: (no message)"
        print(traceback.format_exc(), flush=True)
        raw_lower = raw.lower()
        if any(kw in raw_lower for kw in ("429", "quota", "rate limit", "resource_exhausted")):
            err_msg = "The AI service is currently rate-limited. Please wait a moment and try again."
        elif any(kw in raw_lower for kw in ("500", "internal", "embedding", "overloaded")):
            err_msg = "The AI service returned a temporary error. Please try again in a few seconds."
        elif "expired" in raw_lower or "token" in raw_lower:
            err_msg = "Your session may have expired. Please refresh the page and sign in again."
        else:
            err_msg = "An unexpected error occurred. Please try again."
        return jsonify({'error': err_msg}), 500

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})

# ---------------------- Run App ----------------------
if __name__ == '__main__':
    # debug=False avoids DebuggedApplication multiprocessing PermissionError on /dev/shm (e.g. WSL)
    app.run(debug=False, host='0.0.0.0', port=BACKEND_PORT)
